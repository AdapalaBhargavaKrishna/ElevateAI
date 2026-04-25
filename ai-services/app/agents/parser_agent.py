import pdfplumber
from docx import Document
import io
import re
import logging
from app.agents.llm_service import LLMService

logger = logging.getLogger(__name__)


class ResumeParserAgent:
    SUPPORTED_FORMATS = {"pdf", "docx"}
    EXPECTED_KEYS = [
        "name", "email", "phone", "location", "summary",
        "experience", "projects", "education", "skills",
        "certifications", "languages", "achievements", "coding_profiles"
    ]

    def __init__(self):
        self.llm = LLMService()

    def run(self, file_bytes: bytes = None, filename: str = None, resume_text: str = None) -> dict:
        if resume_text:
            raw_text = self._clean_text(resume_text)
        elif file_bytes and filename:
            print(f"[ParserAgent] Extracting text from '{filename}'...")
            raw_text = self._extract_text(file_bytes, filename)
        else:
            raise ValueError("Provide either resume_text or file_bytes + filename.")

        print(f"[ParserAgent] Extracted {len(raw_text.split())} words. Structuring with LLM...")
        return self._structure_with_llm(raw_text)

    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file type: .{ext}")
        text = self._extract_pdf(file_bytes) if ext == "pdf" else self._extract_docx(file_bytes)
        text = self._clean_text(text)
        if not text.strip():
            raise ValueError("No text could be extracted. File may be image-based.")
        return text

    def _extract_pdf(self, file_bytes: bytes) -> str:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.warning(f"[ParserAgent] Page {i+1} has no extractable text")
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        return text.strip()

    def _structure_with_llm(self, raw_text: str) -> dict:
        prompt = f"""
You are a resume parser. Extract and structure the following resume into JSON with these exact keys:
- name
- email
- phone
- location
- summary
- experience (list of: company, role, duration, responsibilities) — ONLY real work experience/jobs. If none, return null.
- projects (list of: title, description, technologies, link) — ONLY personal/academic projects. If none, return null.
- education (list of: institution, degree, year)
- skills (flat list of strings)
- certifications (list)
- languages (list)
- achievements (list of: title, description, date) — awards, hackathons, competitions. If none, return null.
- coding_profiles (list of: platform, username, link) — LeetCode, GitHub, HackerRank etc. If none, return null.

IMPORTANT RULES:
- Do NOT mix projects into experience.
- Do NOT mix coding profiles into skills.
- For each project, always include a non-empty "title" using the heading/project name from resume text.
- For each project, keep "technologies" as an array of strings.
- If project has a link label like "GitHub" or "Live", still return the URL in "link" when visible.
- Extract all profile/contact links listed in header/footer (GitHub, LinkedIn, LeetCode, portfolio) into coding_profiles.
- If no work experience, set "experience" to null.
- If no projects, set "projects" to null.
- If no achievements, set "achievements" to null.
- If no coding profiles, set "coding_profiles" to null.

Return ONLY valid JSON. No markdown, no backticks, no extra text.

Resume:
{raw_text}
"""
        return self.llm.generate_json(prompt, expected_keys=self.EXPECTED_KEYS)